import re
import shutil
import uuid
import zipfile
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Any

from fastapi import HTTPException, UploadFile


MAX_UPLOAD_FILE_BYTES = 15_000_000
SUPPORTED_EXTENSIONS = {".txt", ".md", ".markdown", ".docx"}
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


@dataclass
class ExtractedDocument:
    filename: str
    ext: str
    text: str
    images: list[dict[str, Any]]
    asset_dir: Path | None = None


async def read_uploaded_document(file: UploadFile, asset_root: Path, static_prefix: str = "/assets") -> ExtractedDocument:
    filename = file.filename or "uploaded.txt"
    ext = Path(filename).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail="Only .txt, .md, .markdown, or .docx files are supported",
        )

    raw = await file.read()
    if not raw:
        raise HTTPException(status_code=400, detail="Uploaded file is empty")
    if len(raw) > MAX_UPLOAD_FILE_BYTES:
        raise HTTPException(status_code=413, detail="File is too large. Please keep it under 15 MB")

    session_id = uuid.uuid4().hex
    asset_dir = asset_root / session_id
    asset_dir.mkdir(parents=True, exist_ok=True)

    if ext == ".docx":
        text, images = extract_docx_text_and_images(raw, asset_dir, static_prefix, session_id)
    else:
        text = decode_text_file(raw)
        images = []

    text = normalize_text(text)
    if len(text.strip()) < 20:
        raise HTTPException(status_code=400, detail="File content is too short to create a presentation")

    return ExtractedDocument(filename=filename, ext=ext, text=text, images=images, asset_dir=asset_dir)


def extract_docx_text_and_images(raw: bytes, asset_dir: Path, static_prefix: str, session_id: str) -> tuple[str, list[dict[str, Any]]]:
    try:
        from docx import Document
    except ImportError:
        raise HTTPException(status_code=500, detail="Missing package: python-docx. Run: pip install -r requirements.txt")

    try:
        document = Document(BytesIO(raw))
    except Exception:
        raise HTTPException(status_code=400, detail="Cannot read .docx file. Make sure it is a valid Word document")

    parts: list[str] = []
    image_refs: list[str] = []

    for idx, para in enumerate(document.paragraphs, start=1):
        text = para.text.strip()
        if text:
            parts.append(text)
        for drawing in para._element.xpath('.//*[local-name()="blip"]'):
            rid = drawing.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
            if rid:
                image_refs.append(rid)
                if text:
                    parts.append(f"[Image near paragraph {idx}: {text[:140]}]")

    for table_index, table in enumerate(document.tables, start=1):
        parts.append(f"[Table {table_index}]")
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    text = "\n".join(parts).strip()
    if not text:
        raise HTTPException(status_code=400, detail="DOCX file does not contain readable text")

    images = extract_images_from_docx_zip(raw, asset_dir, static_prefix, session_id)
    # Preserve first-seen document order when possible.
    rid_to_index = {rid: i for i, rid in enumerate(image_refs)}
    images.sort(key=lambda item: rid_to_index.get(item.get("relationship_id", ""), 9999))
    return text, images


def extract_images_from_docx_zip(raw: bytes, asset_dir: Path, static_prefix: str, session_id: str) -> list[dict[str, Any]]:
    images: list[dict[str, Any]] = []
    media_prefix = "word/media/"
    allowed_suffixes = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tif", ".tiff"}

    try:
        with zipfile.ZipFile(BytesIO(raw)) as zf:
            relationship_by_target = _docx_relationship_targets(zf)
            for index, info in enumerate(zf.infolist(), start=1):
                if not info.filename.startswith(media_prefix):
                    continue
                suffix = Path(info.filename).suffix.lower()
                if suffix not in allowed_suffixes:
                    continue

                data = zf.read(info.filename)
                if not data:
                    continue
                image_id = f"img_{index:03d}"
                out_name = f"{image_id}{'.jpg' if suffix == '.jpeg' else suffix}"
                out_path = asset_dir / out_name
                out_path.write_bytes(data)
                target = info.filename.replace("word/", "")
                rel_id = relationship_by_target.get(target) or relationship_by_target.get("/" + target) or ""
                images.append({
                    "image_id": image_id,
                    "filename": Path(info.filename).name,
                    "path": str(out_path),
                    "url": f"{static_prefix}/{session_id}/{out_name}",
                    "relationship_id": rel_id,
                    "caption": f"Image {len(images) + 1} from the DOCX file",
                })
    except zipfile.BadZipFile:
        raise HTTPException(status_code=400, detail="Invalid .docx file")

    return images[:30]


def _docx_relationship_targets(zf: zipfile.ZipFile) -> dict[str, str]:
    targets: dict[str, str] = {}
    rels_name = "word/_rels/document.xml.rels"
    if rels_name not in zf.namelist():
        return targets
    try:
        import xml.etree.ElementTree as ET
        root = ET.fromstring(zf.read(rels_name))
        ns = {"r": "http://schemas.openxmlformats.org/package/2006/relationships"}
        for rel in root.findall("r:Relationship", ns):
            rel_id = rel.attrib.get("Id", "")
            target = rel.attrib.get("Target", "")
            if rel_id and target:
                targets[target] = rel_id
    except Exception:
        return targets
    return targets


def decode_text_file(raw: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1258", "cp1252"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise HTTPException(status_code=400, detail="Cannot decode text file. Please save it as UTF-8")


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def cleanup_asset_dir(path: Path | None):
    if path and path.exists():
        shutil.rmtree(path, ignore_errors=True)
